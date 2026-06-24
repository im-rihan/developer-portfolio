"""
Generate professional Resume in Word (.docx) and PDF formats.
Run: python generate_resume.py
Output: resume/resume.docx, resume/resume.pdf
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent
OUTPUT_DOCX = SCRIPT_DIR / "resume.docx"
OUTPUT_PDF = SCRIPT_DIR / "resume.pdf"

# ── Resume content (edit personal details here) ──────────────────────────────

PERSONAL = {
    "name": "Rihan Mohammed",
    "title": "Full Stack Developer · HomeAbroad Inc. | Ziffy.ai",
    "email": "im.rihan.dev@gmail.com",
    "phone": "+91 76820 78927",
    "location": "Puri, Odisha, India",
    "address": "Bhagabanpur, Panaspada, Brahmagiri, Puri 752011",
    "linkedin": "linkedin.com/in/im-rihan",
    "github": "github.com/im-rihan",
    "initials": "RM",
}

SUMMARY = (
    "Passionate Full Stack Developer with nearly 4 years at HomeAbroad Inc. and Ziffy.ai, "
    "building production fintech and real-estate platforms at scale. Skilled in JavaScript, "
    "React, Next.js, Node.js, Python, and cloud infrastructure — from AI-native property search "
    "and NestJS APIs to PHP webhook backends, data ingestion pipelines, and AWS deployments. "
    "Committed to delivering innovative, user-focused solutions through effective collaboration "
    "and continuous learning."
)

SKILLS = {
    "Frontend": ["React", "Next.js 15", "TypeScript", "JavaScript", "Tailwind CSS", "Webpack", "Zustand", "TanStack Query", "Formik", "PrimeReact"],
    "Backend": ["NestJS", "Node.js", "PHP 8", "TypeORM", "Express", "REST APIs", "BullMQ", "SSE Streaming"],
    "Data & Search": ["MySQL", "Redis", "Typesense", "Python", "Data Analysis", "Data Ingestion"],
    "AI & Integrations": ["LangChain", "OpenAI", "Google Gemini", "Zoho CRM", "Twilio", "SendGrid"],
    "DevOps & Cloud": ["CI/CD", "GitHub Actions", "Docker", "AWS Lambda", "Amazon EC2", "Vercel", "Hetzner", "Cloudflare", "Nginx"],
}

EXPERIENCE = [
    {
        "title": "Full Stack Engineer",
        "company": "Ziffy.ai",
        "type": "Full-time",
        "dates": "Jan 2025 – Present",
        "duration": "1 yr 6 mos · Remote",
        "subtitle": "AI-native real estate investment platform — ziffy.ai",
        "bullets": [
            "Architected and shipped the Next.js 15 / React 19 frontend for Ziffy.ai with dual-brand support (Ziffy + HomeAbroad), deployed on Vercel with ISR, BunnyCDN sitemaps, and brand-specific analytics.",
            "Built AI-powered property search with Server-Sent Events (SSE) streaming, real-time filter sync via Zustand, and integration with NestJS backend using OpenAI/Perplexity and Typesense full-text search.",
            "Delivered SEO-first programmatic listing pages (city/state/zip catch-all routes), property detail SSR, investment calculators (DSCR, cash flow), and mortgage pre-approval portal with document upload workflows.",
            "Integrated multi-channel analytics (GA4, Google Ads, Facebook Pixel, LogRocket, Vercel Analytics) and Fingerprint.js fraud detection on lead capture forms.",
            "Established dev tooling with Husky, lint-staged, Prettier, and TanStack Query; migrated legacy webpack multi-app frontend to modern App Router architecture.",
        ],
    },
    {
        "title": "Full Stack Developer",
        "company": "HomeAbroad Inc.",
        "type": "Full-time",
        "dates": "Apr 2022 – Present",
        "duration": "4 yrs 3 mos · Remote",
        "subtitle": "Fintech/real-estate platform — software design and full-stack application development",
        "bullets": [
            "Built and maintained the core NestJS REST API (appi) — auth (JWT, OAuth, magic links), property search, loan estimates, pre-approval, CRM sync, document management, and AI features via LangChain/LangGraph with TypeORM/MySQL, Redis throttling, and BullMQ email queues.",
            "Developed the React multi-app frontend (ha-realtor-plat) with Webpack/Tailwind — agent dashboard, investor portal, property maps (Leaflet), loan calculators, term sheets, and document portal with role-based routing for RE agents, MLOs, and clients.",
            "Owned the PHP integration backend (3rdpartycomms) — 60+ webhook receivers, 40+ cron jobs, and 12+ internal agent tools (ClearPath AI, Match AI, Power Dialer, comms centers) integrating Twilio, Zoho CRM, SendGrid, Retell AI, Floify, and Google Gemini.",
            "Authored reusable TypeScript loan estimate library (estimate-calculator) — fees, DSCR, liquidity, Excel rate-sheet parsing; consumed across API and frontend.",
            "Shipped AWS Lambda headless browser service (mortgage-pricer) — Puppeteer scrapers for 11 Non-QM/DSCR lender pricing portals with S3 audit screenshots and API Gateway deployment.",
            "Built property data pipelines — Python scrapers (Zillow, Roofstock, Homes.com, HouseCanary) + Node.js ingestion to MySQL/Typesense with SSH tunnel staging workflows.",
            "Led AWS Lightsail → Hetzner infrastructure migration with xtrabackup streaming, automated bash tooling, and zero-downtime DNS cutover procedures.",
            "Implemented GitHub Actions CI/CD pipelines — build, SFTP upload, SSH deploy to production with staging/main workflows and rollback support.",
            "Maintained WordPress marketing site with custom Tailwind theme, Gutenberg blocks, SEO/lead capture integrated with Zoho CRM and webhook backend.",
        ],
    },
]

PROJECTS = [
    ("Ziffy.ai Platform", "Next.js 15 · React 19 · Zustand · SSE · Vercel", "AI-native investor platform with streaming NLP search, SEO listings, DSCR calculators, and dual-brand deployment."),
    ("appi — Core API", "NestJS · TypeORM · Redis · Typesense · LangChain", "Modular REST API powering property search, loan estimates, auth, CRM sync, and AI-assisted SEO content."),
    ("3rdpartycomms", "PHP 8 · MySQL · Redis · Cloudflare Zero Trust", "Webhook-driven integration hub with agent AI tools, nurture campaigns, and multi-channel comms orchestration."),
    ("mortgage-pricer", "TypeScript · Puppeteer · AWS Lambda · Serverless", "Headless browser microservice scraping live rates from 11 lender portals with pluggable scraper registry."),
    ("estimate-calculator", "TypeScript · Jest · Zero-dep library", "Reusable mortgage calculation engine — fees, liquidity, DSCR, points/pricing scenarios, Excel rate sheets."),
    ("data-pipelines", "Python · Node.js · Typesense · AWS S3", "End-to-end property data acquisition — multi-source scrapers with chunked/resumable ingestion to MySQL."),
]

EDUCATION = [
    {
        "degree": "Post Graduate Diploma in Computer Application (PGDCA)",
        "school": "Utkal University",
        "years": "",
    },
    {
        "degree": "Honors / Regents High School Diploma",
        "school": "HDVSc Degree College, Panaspada, Puri",
        "years": "2016 – 2019",
    },
]

CERTIFICATIONS = [
    "Practical Database Course for Beginners: 6 courses in 1",
    "Build an Amazon Affiliate E-Commerce Store from Scratch",
    "Introduction to Programming Using HTML and CSS",
    "JavaScript Algorithms and Data Structures",
]

# ── Colors ────────────────────────────────────────────────────────────────────

TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_horizontal_line(doc, color="0F766E"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def styled_run(paragraph, text, bold=False, italic=False, size=10, color=None, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, text.upper(), bold=True, size=12, color=TEAL)
    add_horizontal_line(doc)


def generate_docx():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True
    left = header_table.rows[0].cells[0]
    right = header_table.rows[0].cells[1]

    p = left.paragraphs[0]
    styled_run(p, PERSONAL["name"], bold=True, size=22, color=DARK)
    p = left.add_paragraph()
    styled_run(p, PERSONAL["title"], bold=True, size=12, color=TEAL)

    contact_lines = [
        PERSONAL["email"],
        PERSONAL["phone"],
        PERSONAL["location"],
        PERSONAL.get("address", ""),
        PERSONAL["linkedin"],
        PERSONAL["github"],
    ]
    for line in contact_lines:
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        styled_run(p, line, size=9, color=MUTED)

    add_horizontal_line(doc)

    # Summary
    section_heading(doc, "Professional Summary")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED

    # Skills
    section_heading(doc, "Technical Skills")
    for category, items in SKILLS.items():
        p = doc.add_paragraph()
        styled_run(p, f"{category}: ", bold=True, size=10, color=DARK)
        styled_run(p, " · ".join(items), size=9, color=MUTED)

    # Experience
    section_heading(doc, "Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        styled_run(p, job["title"], bold=True, size=11, color=DARK)
        styled_run(p, f"  |  {job['company']} · {job['type']}", size=10, color=TEAL)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        styled_run(p, job["dates"], bold=True, size=9, color=DARK)
        styled_run(p, f"  ·  {job['duration']}", size=9, color=MUTED)

        p = doc.add_paragraph()
        styled_run(p, job["subtitle"], italic=True, size=9, color=MUTED)

        for bullet in job["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK

    # Projects
    section_heading(doc, "Key Projects")
    proj_table = doc.add_table(rows=len(PROJECTS) // 2 + len(PROJECTS) % 2, cols=2)
    proj_table.autofit = True
    for idx, (name, stack, desc) in enumerate(PROJECTS):
        row, col = divmod(idx, 2)
        cell = proj_table.rows[row].cells[col]
        set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        styled_run(p, name, bold=True, size=10, color=TEAL)
        p = cell.add_paragraph()
        styled_run(p, stack, size=8, color=AMBER)
        p = cell.add_paragraph()
        styled_run(p, desc, size=8.5, color=MUTED)

    # Education
    section_heading(doc, "Education")
    for edu in EDUCATION:
        p = doc.add_paragraph()
        styled_run(p, edu["degree"], bold=True, size=10, color=DARK)
        styled_run(p, f"  —  {edu['school']}", size=10, color=TEAL)
        if edu.get("years"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            styled_run(p, edu["years"], size=9, color=MUTED)

    # Certifications
    section_heading(doc, "Certifications")
    for cert in CERTIFICATIONS:
        p = doc.add_paragraph(cert, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK

    doc.save(str(OUTPUT_DOCX))
    print(f"[OK] Word document saved: {OUTPUT_DOCX}")


def generate_pdf():
    from fpdf import FPDF

    class ResumePDF(FPDF):
        def header(self):
            pass

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(100, 116, 139)
            self.cell(0, 8, f"Page {self.page_no()}", align="C")

        def section_title(self, title):
            self.ln(4)
            self.set_font("Helvetica", "B", 11)
            self.set_text_color(15, 118, 110)
            self.cell(0, 7, title.upper(), new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(15, 118, 110)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)

        def bullet(self, text):
            self.set_font("Helvetica", "", 9)
            self.set_text_color(30, 41, 59)
            self.multi_cell(0, 4.5, f"  •  {text}")
            self.ln(1)

    pdf = ResumePDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    # Header bar
    pdf.set_fill_color(15, 23, 42)
    pdf.rect(0, 0, 210, 38, "F")

    pdf.set_xy(15, 10)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 10, PERSONAL["name"], new_x="LMARGIN", new_y="NEXT")

    pdf.set_x(15)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(245, 158, 11)
    pdf.cell(0, 6, PERSONAL["title"], new_x="LMARGIN", new_y="NEXT")

    pdf.set_xy(15, 28)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(203, 213, 225)
    contact = f"{PERSONAL['email']}  |  {PERSONAL['phone']}  |  {PERSONAL['location']}  |  {PERSONAL['linkedin']}"
    pdf.cell(0, 5, contact, new_x="LMARGIN", new_y="NEXT")

    pdf.set_y(44)

    # Summary
    pdf.section_title("Professional Summary")
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(100, 116, 139)
    pdf.multi_cell(0, 5, SUMMARY)
    pdf.ln(2)

    # Skills
    pdf.section_title("Technical Skills")
    for category, items in SKILLS.items():
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(30, 41, 59)
        pdf.cell(35, 5, f"{category}:", new_x="END")
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(100, 116, 139)
        pdf.multi_cell(0, 5, " · ".join(items))
        pdf.ln(1)

    # Experience
    pdf.section_title("Experience")
    for job in EXPERIENCE:
        if pdf.get_y() > 250:
            pdf.add_page()

        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(30, 41, 59)
        pdf.cell(0, 5, job["title"], new_x="END")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(15, 118, 110)
        pdf.cell(0, 5, f"  {job['dates']}", align="R", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(15, 118, 110)
        pdf.cell(0, 5, f"{job['company']} · {job['type']}", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "I", 8.5)
        pdf.set_text_color(100, 116, 139)
        pdf.cell(0, 4, f"{job['subtitle']}  ({job['duration']})", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for bullet in job["bullets"]:
            if pdf.get_y() > 265:
                pdf.add_page()
            pdf.bullet(bullet)
        pdf.ln(3)

    # Projects
    if pdf.get_y() > 220:
        pdf.add_page()
    pdf.section_title("Key Projects")
    for name, stack, desc in PROJECTS:
        if pdf.get_y() > 265:
            pdf.add_page()
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_text_color(15, 118, 110)
        pdf.cell(0, 5, name, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(245, 158, 11)
        pdf.cell(0, 4, stack, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(100, 116, 139)
        pdf.multi_cell(0, 4.5, desc)
        pdf.ln(2)

    # Education
    pdf.section_title("Education")
    for edu in EDUCATION:
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_text_color(30, 41, 59)
        pdf.cell(0, 5, f"{edu['degree']} — {edu['school']}", new_x="LMARGIN", new_y="NEXT")
        if edu.get("years"):
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(100, 116, 139)
            pdf.cell(0, 5, edu["years"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # Certifications
    pdf.section_title("Certifications")
    for cert in CERTIFICATIONS:
        pdf.bullet(cert)

    pdf.output(str(OUTPUT_PDF))
    print(f"[OK] PDF saved: {OUTPUT_PDF}")


def generate_pdf_from_html():
    """Try Playwright HTML→PDF for higher fidelity; fall back to fpdf2."""
    html_path = SCRIPT_DIR / "resume.html"
    if not html_path.exists():
        return False
    try:
        from playwright.sync_api import sync_playwright

        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(html_path.as_uri(), wait_until="networkidle")
            page.pdf(
                path=str(OUTPUT_PDF),
                format="A4",
                print_background=True,
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            )
            browser.close()
        print(f"[OK] PDF (HTML/Playwright) saved: {OUTPUT_PDF}")
        return True
    except Exception as e:
        print(f"  Playwright PDF skipped ({e}), using fpdf2 fallback...")
        return False


if __name__ == "__main__":
    print("Generating resume documents...\n")
    generate_docx()
    if not generate_pdf_from_html():
        generate_pdf()
    print("\nDone! Files are in:", SCRIPT_DIR)
    print("  - resume.html  - open in browser, Print -> Save as PDF")
    print("  - resume.docx  - Microsoft Word")
    print("  - resume.pdf   - PDF format")
